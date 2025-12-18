import os, re, uuid, subprocess
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_file
import pdfplumber
from spellchecker import SpellChecker
from docx import Document
from docx.shared import Pt

# NOTE: This path is specific to a Windows installation of LibreOffice.
# It must be adjusted for Linux/macOS environments (e.g., 'soffice' or '/usr/bin/soffice').
LIBREOFFICE_PATH = r"C:\Program Files\LibreOffice\program\soffice.exe"

app = Flask(__name__)
UPLOAD_FOLDER = Path('uploads')
OUTPUT_FOLDER = Path('outputs')
STATIC_SAMPLE_PATH = Path('static') / 'sample.pdf'
UPLOAD_FOLDER.mkdir(exist_ok=True)
OUTPUT_FOLDER.mkdir(exist_ok=True)

spell = SpellChecker()
sessions = {}

# ------------------ Utility functions ------------------
def clean_text(text):
    """Cleans text by removing PDF artifacts and performing spell correction."""
    text = re.sub(r'\(cid:\d+\)', '', text)
    words = re.split(r'(\w+)', text)
    return ''.join([spell.correction(word) if word.isalpha() else word for word in words])

def identify_fields_to_fill(cleaned_text):
    """Identifies field labels followed by blanks (e.g., Name____)."""
    pattern = r"([A-Za-z ]+?)\s*_{3,}"
    found = re.findall(pattern, cleaned_text)
    out = []
    for f in found:
        f = re.sub(r'\s+', ' ', f.strip())
        if f not in out:
            out.append(f)
    return out

def modify_fields(cleaned_text):
    """Replaces blanks with a standardized 'Label: __' format."""
    return re.sub(r"([A-Za-z ]+?)\s*_{3,}", r"\1: __", cleaned_text)

class AbortProcess(Exception):
    pass

# ✅ ADDED: Helper function to find the field key, ignoring case
def _find_field_key_case_insensitive(fields_list, input_field_name):
    """
    Finds the exact field name in the list, ignoring case.
    Returns the exact key if found, otherwise None.
    """
    input_lower = input_field_name.lower().strip()
    for field in fields_list:
        if field.lower().strip() == input_lower:
            return field
    return None

def validate_input_text(user_input, field_name):
    """Validates user input based on the expected field type."""
    if user_input is None:
        raise ValueError("Empty input")
    if "stop" in user_input.lower():
        raise AbortProcess("User stopped")
    if "skip" in user_input.lower():
        return ""
    if "repeat message" in user_input.lower():
        return "__REPEAT__"

    if field_name.lower().startswith("gender"):
        v = user_input.strip().lower()
        if v in ["female", "f"]:
            return "Female"
        elif v in ["male", "m", "mail"]:
            return "Male"
        else:
            raise ValueError("Enter Male or Female")

    if field_name.lower().startswith("age"):
        nums = re.findall(r'\b([1-9][0-9]?[0-9]?)\b', user_input)
        if nums:
            age = int(nums[0])
            if 1 <= age <= 150:
                return age
        raise ValueError("Enter valid age 1-150")

    if "phone" in field_name.lower() or "number" in field_name.lower():
        num = ''.join(filter(str.isdigit, user_input))
        if len(num) >= 10:
            return num[-10:]
        raise ValueError("Enter 10-digit phone")

    if "email" in field_name.lower():
        if re.match(r'^[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}$', user_input.strip()):
            return user_input.strip().lower()
        else:
            raise ValueError("Enter valid email")

    return user_input.strip().title()

def save_word_from_responses(session_id, responses, font_size=14):
    """Generates a DOCX file containing the collected responses."""
    filename = OUTPUT_FOLDER / f"{session_id}_filled.docx"
    doc = Document()
    for k, v in responses.items():
        p = doc.add_paragraph()
        run = p.add_run(f"{k}: ")
        run.font.size = Pt(font_size)
        value_text = f"___{v}___"
        run = p.add_run(value_text)
        run.font.size = Pt(font_size)
        run.font.underline = True
    doc.save(str(filename))
    return filename

def convert_docx_to_pdf(docx_path):
    """Converts DOCX to PDF using LibreOffice via subprocess."""
    try:
        pdf_path = str(docx_path).replace(".docx", ".pdf")
        subprocess.run(
            [
                LIBREOFFICE_PATH,
                '--headless',
                '--convert-to', 'pdf',
                str(docx_path),
                '--outdir', str(OUTPUT_FOLDER)
            ],
            check=True
        )
        return pdf_path if os.path.exists(pdf_path) else None
    except Exception as e:
        print(f"LibreOffice conversion failed: {e}")
        return None

# ------------------ Flask routes ------------------
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload_pdf', methods=['POST'])
def upload_pdf():
    if 'pdf' not in request.files:
        return jsonify(success=False, error="No file")
    pdf_file = request.files['pdf']
    if pdf_file.filename == '':
        return jsonify(success=False, error="No file")

    uid = uuid.uuid4().hex
    saved_path = UPLOAD_FOLDER / f"{uid}.pdf"
    pdf_file.save(str(saved_path))

    try:
        text_full = ""
        modified_full = ""
        fields = []
        with pdfplumber.open(str(saved_path)) as pdf:
            for p in pdf.pages:
                t = p.extract_text() or ""
                c = clean_text(t)
                text_full += c + "\n"
                fields.extend(identify_fields_to_fill(c))
                modified_full += modify_fields(c) + "\n"
        dedup = []
        for f in fields:
            if f not in dedup:
                dedup.append(f)
        return jsonify(success=True, cleaned_text=text_full.strip(),
                       modified_text=modified_full.strip(), fields=dedup)
    except Exception as e:
        return jsonify(success=False, error=str(e))

@app.route('/use_sample', methods=['POST'])
def use_sample():
    if not STATIC_SAMPLE_PATH.exists():
        return jsonify(success=False, error=f"Sample file not found at {STATIC_SAMPLE_PATH}")
    try:
        text_full = ""
        modified_full = ""
        fields = []
        with pdfplumber.open(str(STATIC_SAMPLE_PATH)) as pdf:
            for p in pdf.pages:
                t = p.extract_text() or ""
                c = clean_text(t)
                text_full += c + "\n"
                fields.extend(identify_fields_to_fill(c))
                modified_full += modify_fields(c) + "\n"
        dedup = []
        for f in fields:
            if f not in dedup:
                dedup.append(f)
        return jsonify(success=True, cleaned_text=text_full.strip(),
                       modified_text=modified_full.strip(), fields=dedup)
    except Exception as e:
        return jsonify(success=False, error=str(e))

@app.route('/start_session', methods=['POST'])
def start_session():
    data = request.get_json() or {}
    fields = data.get('fields') or []
    if not fields:
        return jsonify(success=False, error="No fields")
    session_id = uuid.uuid4().hex
    sessions[session_id] = {
        'fields': fields,
        'index': 0,
        'responses': {},
        'mode': 'filling'  # 'filling' | 'review'
    }
    return jsonify(success=True, session_id=session_id, fields=fields)

@app.route('/process_input', methods=['POST'])
def process_input():
    data = request.get_json() or {}
    session_id = data.get('session_id')
    answer = data.get('answer', '')
    if not session_id or session_id not in sessions:
        return jsonify(success=False, error="Invalid session")
    session = sessions[session_id]
    idx = session['index']
    fields = session['fields']
    if idx >= len(fields):
        # Should be handled by client logic calling /validate_field instead
        return jsonify(success=False, error="All initial filling complete. Please use /validate_field for corrections.")
    current_field = fields[idx]
    try:
        validated = validate_input_text(answer, current_field)
        if validated == "__REPEAT__":
            return jsonify(success=True, message="Repeating instructions.", next_field=current_field, next_index=idx)
        session['responses'][current_field] = validated
        session['index'] = idx + 1
        if session['index'] >= len(fields):
            session['mode'] = 'review'
            return jsonify(
                success=True,
                completed=True,
                responses=session['responses'], # Return all responses for review display
                message="All fields filled. Please review and confirm the details.",
                validated_value=validated
            )
        next_field = fields[session['index']]
        return jsonify(
            success=True,
            message=f"Recorded {current_field}. Next: {next_field}",
            next_field=next_field,
            next_index=session['index'],
            validated_value=validated
        )
    except AbortProcess as ap:
        sessions.pop(session_id, None)
        return jsonify(success=True, message=str(ap), completed=True)
    except ValueError as ve:
        return jsonify(success=False, error=str(ve))

@app.route('/get_session_state', methods=['POST'])
def get_session_state():
    """Return current fields, index, responses, and mode."""
    data = request.get_json() or {}
    session_id = data.get('session_id')
    if not session_id or session_id not in sessions:
        return jsonify(success=False, error="Invalid session")
    s = sessions[session_id]
    return jsonify(
        success=True,
        fields=s['fields'],
        index=s['index'],
        responses=s['responses'],
        mode=s.get('mode', 'filling')
    )

@app.route('/validate_field', methods=['POST'])
def validate_field():
    """
    Validate and update a single field during review/correction.
    Payload: { session_id, field_name, value }
    """
    data = request.get_json() or {}
    session_id = data.get('session_id')
    field_name = data.get('field_name', '')
    value = data.get('value', '')
    if not session_id or session_id not in sessions:
        return jsonify(success=False, error="Invalid session")
    session = sessions[session_id]

    # ✅ Uses the new helper function
    exact_key = _find_field_key_case_insensitive(session['fields'], field_name)
    if not exact_key:
        return jsonify(success=False, error=f"Field '{field_name}' not found. Choose one of: {', '.join(session['fields'])}")

    try:
        validated = validate_input_text(value, exact_key)
        if validated == "__REPEAT__":
            return jsonify(success=False, error="Repeat is not allowed during correction.")

        session['responses'][exact_key] = validated
        session['mode'] = 'review' # Keep mode as review after correction
        return jsonify(success=True, field=exact_key, validated_value=validated, responses=session['responses'])
    except ValueError as ve:
        return jsonify(success=False, error=str(ve))
    except AbortProcess as ap:
        sessions.pop(session_id, None)
        return jsonify(success=True, message=str(ap), completed=True)

@app.route('/finalize_form', methods=['POST'])
def finalize_form():
    """Finalizes the form by generating DOCX and PDF files."""
    data = request.get_json() or {}
    session_id = data.get('session_id')
    if not session_id or session_id not in sessions:
        return jsonify(success=False, error="Invalid session")
    session = sessions[session_id]
    
    if session.get('mode') != 'review':
         return jsonify(success=False, error="Cannot finalize: form is not yet complete or reviewed.")

    responses = session['responses']
    docx_file = save_word_from_responses(session_id, responses)
    pdf_file = convert_docx_to_pdf(docx_file)
    pdf_url_result = f'/download/{Path(pdf_file).name}' if pdf_file and Path(pdf_file).exists() else ""
    
    # Clean up session after final generation
    sessions.pop(session_id, None)
    
    return jsonify(success=True,
                   docx_url=f'/download/{docx_file.name}',
                   pdf_url=pdf_url_result,
                   message="Form finalized and ready for download.")

@app.route('/download/<filename>')
def download_file(filename):
    fpath = OUTPUT_FOLDER / filename
    if fpath.exists():
        if filename.endswith('.pdf'):
            mimetype = 'application/pdf'
        elif filename.endswith('.docx'):
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            mimetype = 'application/octet-stream'
        return send_file(str(fpath), mimetype=mimetype, as_attachment=True, download_name=filename)
    return "File not found", 404

if __name__ == '__main__':
    app.run(debug=True)
